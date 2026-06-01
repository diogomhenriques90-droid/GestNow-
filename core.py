"""
GESTNOW v3 - core.py
Funções base, utilitários, segurança, UI components, notificações, offline, QR Code, SMTP

PERFORMANCE FIXES v3.3:
  1. inv() selectivo — versioning via _fv em session_state, só limpa o necessário
  2. _cached_load_db() recebe _v → cache miss só no ficheiro alterado
  3. _load_users_cached() com @st.cache_data(ttl=1800) — evita 4+ leituras GCS/rerun
  4. check_connection_status() com @st.cache_data(ttl=30) — evita chamada GCS desnecessária
  5. _gcs_client() com @st.cache_resource — cliente GCS criado uma vez por instância
  6. _cached_load_db() TTL aumentado 300s→3600s (inv() garante frescura nas escritas)
"""
import streamlit as st
import pandas as pd
import os, re, secrets, io, base64, bcrypt, logging, uuid, hashlib, json, time
from datetime import datetime, timedelta, date
from google.cloud import storage as gcs
import plotly.express as px
from streamlit_folium import folium_static
import folium

# =============================================================================
# DESIGN SYSTEM
# =============================================================================
COLORS = {
    "primary":       "#0F172A",
    "primary_light": "#1E293B",
    "accent":        "#3B82F6",
    "accent_hover":  "#60A5FA",
    "success":       "#10B981",
    "warning":       "#F59E0B",
    "error":         "#EF4444",
    "info":          "#8B5CF6",
    "bg_glass":      "rgba(255,255,255,0.1)",
    "border_glass":  "rgba(255,255,255,0.2)",
    "text_primary":  "#F8FAFC",
    "text_secondary":"#94A3B8",
}

ICONS = {
    "app": "🎛️", "login": "🔐", "admin": "⚡", "technician": "👨‍🔧",
    "dashboard": "📈", "instrumentation": "🧪", "voice": "🎤", "safety": "🛡️",
    "profile": "👤", "reports": "📊", "handover": "✅", "gps": "📍",
    "calibration": "🔬", "material": "📦", "pending": "⏳", "logout": "🚪",
    "save": "💾", "edit": "✏️", "delete": "🗑️", "add": "➕",
    "search": "🔍", "filter": "🔽", "download": "📥", "upload": "📤",
    "check": "✅", "close": "❌", "warning": "⚠️", "info": "ℹ️",
    "calendar": "📅", "clock": "⏰", "user": "👤", "users": "👥",
    "settings": "⚙️", "home": "🏠", "work": "🏗️", "tools": "🔧",
    "equipment": "🔩", "document": "📄", "documents": "📁",
    "chart": "📊", "graph": "📈", "email": "📧", "phone": "📞",
    "location": "📍", "time": "⏱️", "approved": "✅", "rejected": "❌",
    "pending_approval": "⏳",
}

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# =============================================================================
# CONFIGURAÇÕES DE STORAGE
# =============================================================================
GCS_BUCKET = os.environ.get("GCS_BUCKET", "gestnow-dados")

@st.cache_resource
def _gcs_client():
    try:
        return gcs.Client()
    except Exception as e:
        logger.warning(f"GCS client fallback: {e}")
        return None

def _gcs_read(fn):
    try:
        client = _gcs_client()
        if client:
            bucket = client.bucket(GCS_BUCKET)
            blob   = bucket.blob(f"data/{fn}")
            if blob.exists():
                return io.BytesIO(blob.download_as_bytes())
    except Exception as e:
        logger.debug(f"GCS read fallback para {fn}: {e}")
    return None

def _gcs_write(fn, content_bytes):
    try:
        client = _gcs_client()
        if not client:
            logger.error(f"❌ _gcs_write({fn}): cliente GCS não inicializado (credenciais em falta?)")
            st.toast("⚠️ Sem ligação ao GCS — dados não guardados", icon="⚙️")
            return False
        bucket = client.bucket(GCS_BUCKET)
        blob   = bucket.blob(f"data/{fn}")
        blob.metadata = {
            "last_updated": datetime.now().isoformat(),
            "app_version":  "GESTNOW-v3.0"
        }
        blob.upload_from_string(content_bytes, content_type="text/csv")
        return True
    except Exception as e:
        logger.error(f"❌ Erro crítico GCS write {fn}: {e}")
        st.toast("⚠️ Erro ao guardar dados", icon="⚙️")
    return False

# =============================================================================
# GCS BINÁRIO, BACKUP E CONTRATO
# =============================================================================

def _gcs_append_row(fn: str, row: dict) -> bool:
    """
    Appends one CSV row to a GCS file via server-side Compose — zero GCS reads.
    Writes a tiny delta blob, then merges it into the main file.
    Falls back to load_db→save_db on error.
    """
    client = _gcs_client()
    if client is None:
        return False
    try:
        bucket    = client.bucket(GCS_BUCKET)
        main_blob = bucket.blob(f"data/{fn}")
        row_csv   = pd.DataFrame([row]).to_csv(index=False, header=False)
        if not main_blob.exists():
            header = ",".join(str(k) for k in row.keys())
            main_blob.upload_from_string(
                f"{header}\n{row_csv}".encode('utf-8-sig'), content_type="text/csv")
            return True
        delta = bucket.blob(f"data/_d_{uuid.uuid4().hex[:8]}_{fn}")
        delta.upload_from_string(row_csv.encode('utf-8'), content_type="text/csv")
        try:
            main_blob.compose([main_blob, delta])
        finally:
            try:    delta.delete()
            except: pass
        return True
    except Exception as e:
        logger.warning(f"_gcs_append_row fallback {fn}: {e}")
        try:
            df = load_db(fn, list(row.keys()))
            save_db(pd.concat([df, pd.DataFrame([row])], ignore_index=True), fn)
            return True
        except Exception as e2:
            logger.error(f"_gcs_append_row fallback falhou {fn}: {e2}")
            return False


def _gcs_write_binary(data: bytes, filename: str) -> bool:
    try:
        client = _gcs_client()
        if not client:
            logger.error(f"❌ _gcs_write_binary({filename}): cliente GCS não inicializado")
            return False
        bucket = client.bucket(GCS_BUCKET)
        blob   = bucket.blob(f"data/{filename}")
        blob.upload_from_string(data)
        return True
    except Exception as e:
        logger.error(f"❌ Erro GCS write binary {filename}: {e}")
    return False


def _gcs_read_binary(filename: str):
    try:
        client = _gcs_client()
        if client:
            bucket = client.bucket(GCS_BUCKET)
            blob   = bucket.blob(f"data/{filename}")
            if blob.exists():
                return blob.download_as_bytes()
    except Exception as e:
        logger.error(f"❌ Erro GCS read binary {filename}: {e}")
    return None


def _verificar_alerta_backup() -> tuple:
    try:
        buf = _gcs_read("backup_status.json")
        if not buf:
            return 'nunca', None
        data       = json.loads(buf.read().decode('utf-8'))
        ultima_str = data.get('Data_Backup', '')
        if not ultima_str:
            return 'nunca', None
        ultima  = datetime.strptime(ultima_str, "%d/%m/%Y %H:%M")
        diff_h  = (datetime.now() - ultima).total_seconds() / 3600
        if diff_h < 24:
            return 'ok', ultima
        elif diff_h < 48:
            return 'aviso', ultima
        else:
            return 'critico', ultima
    except:
        return 'nunca', None


def _registar_backup(admin_nome: str):
    payload = {
        "Data_Backup": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "Admin": admin_nome
    }
    _gcs_write_binary(json.dumps(payload).encode('utf-8'), "backup_status.json")


def _fill_contrato_template(substituicoes: dict):
    template_bytes = _gcs_read_binary("contrato_template.docx")
    if not template_bytes:
        return None
    try:
        from docx import Document
        from io import BytesIO as _BIO

        doc = Document(_BIO(template_bytes))

        def _sub(para):
            if not any(k in para.text for k in substituicoes):
                return
            for run in para.runs:
                for k, v in substituicoes.items():
                    if k in run.text:
                        run.text = run.text.replace(k, str(v))
            if any(k in para.text for k in substituicoes):
                if para.runs:
                    r0    = para.runs[0]
                    bold  = r0.bold
                    italic= r0.italic
                    fname = r0.font.name
                    fsize = r0.font.size
                    texto = para.text
                    for k, v in substituicoes.items():
                        texto = texto.replace(k, str(v))
                    for run in para.runs:
                        run.text = ""
                    para.runs[0].text   = texto
                    para.runs[0].bold   = bold
                    para.runs[0].italic = italic
                    if fname: para.runs[0].font.name = fname
                    if fsize: para.runs[0].font.size = fsize

        for para in doc.paragraphs:
            _sub(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _sub(para)

        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return out.getvalue()
    except Exception as e:
        logger.error(f"Erro ao preencher contrato: {e}")
        return None

# =============================================================================
# GESTÃO DE DADOS
# =============================================================================
@st.cache_data(ttl=3600, show_spinner=False)
def _cached_load_db(fn, cols_tuple, silent=False, _v=0):
    """Backend cacheado — não chamar directamente, usar load_db()."""
    buf = _gcs_read(fn)
    if buf:
        try:
            df = pd.read_csv(buf, dtype=str, on_bad_lines='skip', encoding='utf-8-sig')
            df.columns = df.columns.str.strip()
            for c in cols_tuple:
                if c.strip() not in df.columns:
                    df[c.strip()] = ""
            return df[[c.strip() for c in cols_tuple]].fillna("")
        except Exception as e:
            if not silent:
                logger.warning(f"⚠️ Fallback CSV local para {fn}: {e}")
            return pd.DataFrame(columns=[c.strip() for c in cols_tuple])
    return pd.DataFrame(columns=[c.strip() for c in cols_tuple])


def load_db(fn, cols, silent=False):
    """Carrega CSV do GCS com invalidação selectiva automática por ficheiro."""
    v = st.session_state.get('_fv', {}).get(fn, 0)
    df = _cached_load_db(fn, tuple(cols), silent=silent, _v=v)
    if fn in _CRITICAL_FILES and not df.empty:
        st.session_state.setdefault('_snap', {})[fn] = df
    return df

# ── Ficheiros críticos — protegidos contra perda de dados ───────────────────
_CRITICAL_FILES = {"registos.csv", "usuarios.csv", "folhas_ponto.csv", "contratos.csv"}


def save_db(df, fn):
    """
    Guarda DataFrame no GCS de forma SEGURA.
    Para ficheiros críticos: verifica perda de registos e cria backup diário.
    Usa snapshot em session_state para evitar leitura GCS antes de cada escrita.
    """
    # ── Segurança para ficheiros críticos ────────────────────────────────────
    if fn in _CRITICAL_FILES:
        try:
            snap = st.session_state.get('_snap', {}).get(fn)
            if snap is not None:
                df_atual = snap
            else:
                buf_atual = _gcs_read(fn)
                df_atual  = pd.read_csv(buf_atual, dtype=str,
                                        on_bad_lines='skip', encoding='utf-8-sig') \
                            if buf_atual else pd.DataFrame()
            if not df_atual.empty:
                n_atual = len(df_atual)
                n_novo  = len(df)
                _criar_backup_diario(fn, df_atual)
                if n_atual > 5 and n_novo < int(n_atual * 0.90):
                    logger.error(
                        f"🚨 BLOQUEADO save_db({fn}): {n_novo} registos "
                        f"vs {n_atual} existentes — perda de "
                        f"{n_atual - n_novo} registos (>10%). Operação cancelada."
                    )
                    return False
        except Exception as e:
            logger.warning(f"Verificação segurança {fn}: {e}")

    # ── Guardar no GCS ────────────────────────────────────────────────────────
    buf = io.StringIO()
    df.to_csv(buf, index=False, encoding='utf-8-sig')
    result = _gcs_write(fn, buf.getvalue().encode('utf-8-sig'))
    if result and fn in _CRITICAL_FILES:
        st.session_state.setdefault('_snap', {})[fn] = df
    return result


def _criar_backup_diario(fn, df_atual):
    """Backup automático diário em backups/YYYY-MM-DD/ficheiro.csv"""
    try:
        today = datetime.now().strftime("%Y-%m-%d")
        cache_key = f"_backup_done_{today}_{fn}"
        if st.session_state.get(cache_key):
            return
        client = _gcs_client()
        if client:
            bucket = client.bucket(GCS_BUCKET)
            blob   = bucket.blob(f"data/backups/{today}/{fn}")
            if not blob.exists():
                buf = io.StringIO()
                df_atual.to_csv(buf, index=False, encoding='utf-8-sig')
                blob.upload_from_string(
                    buf.getvalue().encode('utf-8-sig'), content_type="text/csv")
                logger.info(f"✅ Backup diário: backups/{today}/{fn} ({len(df_atual)} registos)")
            st.session_state[cache_key] = True
    except Exception as e:
        logger.warning(f"Backup diário falhou {fn}: {e}")


def list_backups(fn="registos.csv"):
    """Lista backups disponíveis. Retorna [{date, path, size_kb}]"""
    backups = []
    try:
        client = _gcs_client()
        if client:
            bucket = client.bucket(GCS_BUCKET)
            for blob in bucket.list_blobs(prefix="data/backups/"):
                if blob.name.endswith(fn):
                    parts = blob.name.split('/')
                    if len(parts) >= 4:
                        backups.append({
                            "date":    parts[2],
                            "path":    blob.name,
                            "size_kb": (blob.size or 0) // 1024
                        })
            backups.sort(key=lambda x: x["date"], reverse=True)
    except Exception as e:
        logger.warning(f"list_backups: {e}")
    return backups


def restore_backup(backup_path):
    """Restaura ficheiro a partir de backup. Retorna True se OK."""
    try:
        client = _gcs_client()
        if client:
            bucket = client.bucket(GCS_BUCKET)
            blob_s = bucket.blob(backup_path)
            if blob_s.exists():
                fn_dest   = backup_path.split('/')[-1]
                blob_dest = bucket.blob(f"data/{fn_dest}")
                blob_dest.upload_from_string(
                    blob_s.download_as_bytes(), content_type="text/csv")
                inv(fn_dest)
                logger.info(f"✅ Restaurado: {backup_path} → data/{fn_dest}")
                return True
    except Exception as e:
        logger.error(f"restore_backup: {e}")
    return False

# ─────────────────────────────────────────────────────────────────────────────
# FIX 1 — _load_users_cached com cache TTL=60s
# Substitui todas as chamadas directas ao GCS para ler usuarios.csv.
# Evita 4+ leituras de rede por rerun.
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(ttl=1800, show_spinner=False)
def _load_users_cached():
    """
    Lê usuarios.csv do GCS com cache de 30 minutos (1800s).
    Usar em todo o código em vez de _gcs_read("usuarios.csv") directo.
    Para forçar leitura fresca: _load_users_cached.clear()
    """
    try:
        buf = _gcs_read("usuarios.csv")
        if buf:
            df = pd.read_csv(buf, dtype=str, on_bad_lines='skip',
                             encoding='utf-8-sig')
            df.columns = df.columns.str.strip()
            for col in df.select_dtypes(include='object').columns:
                df[col] = df[col].str.strip()
            return df.fillna("")
    except Exception as e:
        logger.warning(f"_load_users_cached erro: {e}")
    return pd.DataFrame()

# ─────────────────────────────────────────────────────────────────────────────
# FIX 2 — inv() verdadeiramente selectivo via versão por ficheiro
# inv("ficheiro.csv") → incrementa versão só desse ficheiro em session_state
#                       → _cached_load_db recebe _v diferente → cache miss apenas para esse ficheiro
# inv()               → nuclear (st.cache_data.clear()) — só para reset/admin global
# ─────────────────────────────────────────────────────────────────────────────
def inv(ficheiro=None):
    """
    Invalida cache de dados de forma selectiva.

    inv()                → limpa TUDO (operações raras: reset, admin global)
    inv("registos.csv")  → invalida só registos.csv (os outros CSVs ficam cacheados)
    inv("usuarios.csv")  → invalida usuarios.csv + _load_users_cached
    """
    if ficheiro is None:
        # Nuclear — limpa toda a cache do Streamlit
        st.cache_data.clear()
    else:
        # Timestamp ms como _v — chave única por evento, sem colisão entre sessões
        fv = dict(st.session_state.get('_fv', {}))
        fv[ficheiro] = int(time.time() * 1000)
        st.session_state['_fv'] = fv
        if ficheiro == "usuarios.csv":
            _load_users_cached.clear()


def load_all():
    users = load_db("usuarios.csv", [
        "Nome", "Password", "Tipo", "Email", "Telefone", "Cargo",
        "NIF", "NISS", "CC", "CC_Validade", "DataNasc", "Nacionalidade",
        "Morada", "Localidade", "Concelho", "Codigo_Postal", "Naturalidade",
        "Estado_Civil", "Sexo", "Dependentes", "Profissao",
        "Categoria_Profissional", "Habilitacoes_Literarias",
        "Contacto_Emergencia", "Nome_Emergencia", "Grau_Parentesco",
        "Banco_IBAN", "Observacoes", "Tamanho_Camisola", "Tamanho_Calca",
        "Tamanho_Botas", "Local", "Foto",
        "PrecoHora", "PrecoHoraStatus", "PrecoHoraData", "PIN",
        "Campos_Bloqueados", "PDFs_Vistos", "PDFs_Validados",
        "PDFs_Validacao_Data",
        "Perfil_Completo", "Perfil_Data",
        "IBAN_Comprovativo_b64", "IBAN_Data_Upload",
        "Local_Obra", "Cliente_Obra",
        "Contrato_Gerado", "Contrato_Data", "Contrato_b64",
        "Contrato_Enviado", "Contrato_Enviado_Data",
        "Contrato_Assinado", "Contrato_Assinatura_b64", "Contrato_Assinatura_Data",
        "Contrato_Validado_Admin", "Contrato_Validado_Data",
        "Contrato_Local_Obra", "Contrato_Cliente_Obra"
    ])

    obras = load_db("obras_lista.csv", [
        "Obra", "Codigo", "Cliente", "Local", "Ativa",
        "Latitude", "Longitude", "Raio_Validacao",
        "DataInicio", "DataFim", "TipoObra", "AssinaturaObrigatoria", "Logo_b64"
    ])

    frentes = load_db("frentes_lista.csv", [
        "Obra", "Frente", "Tipo", "Responsavel"
    ])

    regs = load_db("registos.csv", [
        "ID", "Data", "Técnico", "Obra", "Frente", "TipoFrente",
        "Turnos", "Relatorio", "Status", "Horas_Total",
        "Localizacao_Checkin", "Localizacao_Checkout", "Periodo"
    ])
    if not regs.empty:
        regs['Data']        = pd.to_datetime(regs['Data'], dayfirst=True, errors='coerce')
        regs['Horas_Total'] = pd.to_numeric(regs['Horas_Total'], errors='coerce').fillna(0)

    fats = load_db("faturas.csv",   ["Numero","Cliente","Valor","Status","Data_Emissao","Obra"])
    docs = load_db("documentos.csv",["Utilizador","Tipo","Nome","Validade"])

    incs = load_db("incidentes.csv", [
        "ID", "Data", "Utilizador", "Solicitante", "Obra", "Tipo",
        "Descricao", "Gravidade", "Status", "Equipamento",
        "Urgencia", "Valor_Estimado", "Fatura_b64",
        "Data_Validacao", "Validado_Por"
    ])

    sw  = load_db("safety_walks.csv",  [
        "Data","Utilizador","Obra","Categoria","Descricao","AcaoCorretiva","Status","Urgencia"
    ])
    obs = load_db("obs_seguranca.csv", [
        "Data","Utilizador","Obra","Tipo","Descricao","Status"
    ])
    equip = load_db("equipamentos.csv",[
        "Obra","Tipo","Descricao","NumSerie","Utilizador","Validade","Estado"
    ])

    diags   = load_db("dialogos.csv",      ["Titulo","Descricao","Tipo","DataCriacao","Atribuidos","Estado"])
    diags_u = load_db("dialogos_users.csv",["Dialogo","Utilizador","DataLeitura","Confirmado"])

    folhas = load_db("folhas_ponto.csv", [
        "ID", "Obra", "Periodo", "Responsavel", "Data_Assinatura",
        "Assinatura_b64", "Selo", "Status",
        "Data", "ChefEquipa", "TotalHoras", "AssinadoCliente", "PDF_b64"
    ])

    comuns   = load_db("comunicados.csv",      ["ID","Titulo","Conteudo","Tipo","Destino","Urgente","Validade"])
    comuns_u = load_db("comunicados_lidos.csv",["ComunicadoID","Utilizador","DataLeitura"])

    req_fer = load_db("req_ferramentas.csv", [
        "ID", "Data", "Solicitante", "Obra", "Descricao",
        "Urgencia", "Foto_b64", "Status", "Data_Validacao", "Validado_Por"
    ])
    req_mat = load_db("req_materiais.csv", [
        "ID", "Data", "Solicitante", "Obra", "Descricao", "Quantidade",
        "Unidade", "Urgencia", "Status", "Tipo", "Litros", "Valor",
        "Data_Abastecimento", "Recibo_b64", "Equipamento", "Valor_Estimado",
        "Fatura_b64", "Data_Validacao", "Validado_Por"
    ])
    req_epi = load_db("req_epis.csv", [
        "ID", "Data", "Solicitante", "Obra", "Item", "Tamanho",
        "Quantidade", "Descricao", "Status", "Data_Validacao", "Validado_Por"
    ])

    avals = load_db("avaliacoes.csv", [
        "Data", "Trabalhador", "Nota_Tecnica", "Nota_Pontualidade",
        "Nota_Trabalho_Eq", "Nota_Proatividade", "Nota_Comunicacao",
        "Media", "Comentarios"
    ])

    inst_acessos = load_db("inst_acessos.csv", [
        "Obra", "Utilizador", "Cargo", "Ativo"
    ])

    diarias_config = load_db("diarias_config.csv", [
        "Obra", "Valor_Diaria", "Atualizado_Em", "Atualizado_Por"
    ])

    diarias_faltas = load_db("diarias_faltas.csv", [
        "ID", "Data", "Técnico", "Obra", "Motivo",
        "Registado_Por", "Registado_Em"
    ])

    folhas_ocr = load_db("folhas_ocr.csv", [
        "ID", "Obra", "Periodo", "Semana_Inicio", "Semana_Fim",
        "Tecnico", "Horas_Folha", "Dias", "Extraido_Em",
        "Extraido_Por", "Imagem_b64", "Confianca"
    ])

    diarias_pagamentos = load_db("diarias_pagamentos.csv", [
        "ID", "Semana_Inicio", "Semana_Fim", "Técnico",
        "Obras", "Dias_Total", "Valor_Total", "IBAN",
        "Status", "Data_Pagamento", "Pago_Por", "Recibo_b64"
    ])

    return (users, obras, frentes, regs, fats, docs, incs, sw, obs, equip,
            diags, diags_u, folhas, comuns, comuns_u, req_fer, req_mat, req_epi,
            avals, inst_acessos, diarias_config, diarias_faltas, diarias_pagamentos,
            folhas_ocr)

# =============================================================================
# UTILITÁRIOS
# =============================================================================
def fh(h):
    if h is None or (isinstance(h, float) and pd.isna(h)):
        return "0h00m"
    try:
        h_float = float(h)
        return f"{int(h_float)}h{int((h_float - int(h_float)) * 60):02d}m"
    except:
        return "0h00m"

def sl(s):
    mapping = {
        "0":  ("Pendente",    "status-pending",    "⏳", COLORS["warning"]),
        "1":  ("Material OK", "status-ok",         "📦✅", COLORS["success"]),
        "2":  ("Calibrado",   "status-calibrated", "🧪", COLORS["info"]),
        "3":  ("Instalado",   "status-installed",  "📍", COLORS["accent"]),
        "4":  ("Concluído",   "status-completed",  "✅🎯", COLORS["success"]),
        "-1": ("Rejeitado",   "status-rejected",   "❌", COLORS["error"]),
    }
    key = str(s).strip() if s else "0"
    return mapping.get(key, ("Desconhecido", "status-unknown", "❓", COLORS["text_secondary"]))

def process_and_compress_image(image_file, max_size=(1280, 1280), quality=85):
    try:
        from PIL import Image
        img = Image.open(image_file)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img.thumbnail(max_size, Image.Resampling.LANCZOS)
        output = io.BytesIO()
        img.save(output, format="JPEG", quality=quality, optimize=True, progressive=True)
        return base64.b64encode(output.getvalue()).decode('utf-8')
    except Exception as e:
        logger.error(f"Erro ao processar imagem: {e}")
        return None

def canvas_to_b64(image_data):
    try:
        from PIL import Image
        img = Image.fromarray(image_data.astype("uint8"), "RGBA")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return base64.b64encode(buf.getvalue()).decode()
    except Exception as e:
        logger.error(f"Erro canvas_to_b64: {e}")
        return None

# =============================================================================
# PWA & METADATA
# =============================================================================
def inject_pwa_meta():
    st.markdown("""
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
    <meta name="theme-color" content="#0F172A">
    <meta name="description" content="GESTNOW - Gestão de Instrumentação Industrial">
    <link rel="manifest" href="/manifest.json">
    <link rel="icon" type="image/png" href="/favicon-32x32.png">
    """, unsafe_allow_html=True)

# =============================================================================
# SEGURANÇA
# =============================================================================
def init_session():
    defaults = {
        'user': None, 'tipo': None, 'cargo': None,
        'data_consulta': date.today(), 'login_attempts': 0,
        'last_activity': datetime.now(), 'session_token': None,
        'language': 'pt', 'obra_ativa': None, 'menu_selected': '',
        '_menu_locked': False
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

def check_timeout():
    if st.session_state.get('user') and st.session_state.get('last_activity'):
        if isinstance(st.session_state['last_activity'], datetime):
            inactive = (datetime.now() - st.session_state['last_activity']).seconds / 60
            if inactive > 120:
                logger.info(f"🔒 Timeout: {st.session_state.get('user')}")
                st.session_state.clear()
                st.toast("🔒 Sessão expirada", icon="🔐")
                st.rerun()
            st.session_state['last_activity'] = datetime.now()

def hp(p):
    return bcrypt.hashpw(p.encode('utf-8'), bcrypt.gensalt(rounds=12)).decode('utf-8')

def cp(p, h):
    try:
        return bcrypt.checkpw(p.encode('utf-8'), h.encode('utf-8'))
    except Exception as e:
        logger.warning(f"⚠️ Erro na verificação: {e}")
        return False

# =============================================================================
# COMPONENTES UI
# =============================================================================
def render_metric(icon, val, lbl, color=None):
    if color is None:
        color = COLORS["accent"]
    st.markdown(f"""
    <div style="text-align:center;padding:16px 20px;
        background:linear-gradient(135deg,{COLORS['primary_light']},{COLORS['primary']});
        border-radius:16px;border:1px solid {COLORS['border_glass']};
        box-shadow:0 4px 20px rgba(0,0,0,0.2);">
        <div style="font-size:2rem;font-weight:800;color:{color};margin-bottom:4px;">{icon}</div>
        <div style="font-size:1.5rem;font-weight:700;color:{COLORS['text_primary']};">{val}</div>
        <div style="font-size:0.85rem;color:{COLORS['text_secondary']};margin-top:4px;">{lbl}</div>
    </div>
    """, unsafe_allow_html=True)

# =============================================================================
# CONSTANTES
# =============================================================================
TIPOS_FRENTE = [
    "Montagem", "Calibração", "Comissionamento", "Testes FAT/SAT",
    "Manutenção", "Troubleshooting", "Outro"
]
CARGOS = [
    "Instrumentista Senior", "Técnico de Campo",
    "Engenheiro de Instrumentação", "Chefe de Equipa", "QA/QC", "Admin"
]
CATEGORIAS_SAFETY_WALK = [
    "EPI Industrial", "Lockout/Tagout", "Espaços Confinados",
    "Trabalho em Altura", "Elétrica", "Pressão", "Outro"
]
REGRAS_OURO = [
    ("🛡️", "EPI Industrial Obrigatório",  "Capacete, óculos, luvas e calçado de segurança."),
    ("⚡", "LOTO - Lockout/Tagout",        "Bloqueio e etiquetagem de energias."),
    ("🪜", "Trabalho em Altura",           "Arnés e linha de vida acima de 1.8m."),
    ("⚡", "Energias Perigosas",           "Verificar ausência de tensão."),
    ("🧪", "Calibração Certificada",       "Usar equipamentos com certificado válido."),
    ("📍", "Procedimentos de Campo",       "Seguir ITRs e checklists."),
    ("🔒", "Acesso Restrito",              "Áreas de instrumentação controladas."),
    ("📋", "Análise de Risco",             "JSA/JHA obrigatório."),
    ("🧤", "Mãos Limpas",                  "Luvas adequadas para instrumentos."),
    ("📱", "Zona Livre de Telemóvel",       "Dispositivos proibidos em áreas classificadas."),
]

# =============================================================================
# CSS GLOBAL
# =============================================================================
GLOBAL_CSS = """
:root {
    --primary: #0F172A; --primary-light: #1E293B;
    --accent: #3B82F6; --accent-hover: #60A5FA;
    --success: #10B981; --warning: #F59E0B;
    --error: #EF4444; --info: #8B5CF6;
    --text-primary: #FFFFFF; --text-secondary: #94A3B8;
    --text-dark: #1E293B; --text-light: #F8FAFC;
    --bg-white: #FFFFFF; --bg-light: #F8FAFC; --bg-dark: #0F172A;
}
.stApp {
    background: linear-gradient(135deg, var(--primary) 0%, #1a1a2e 100%);
    color: var(--text-primary);
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
}
.stTextInput > div > div > input,
.stNumberInput > div > div > input,
.stTextArea > div > div > textarea {
    background: var(--bg-white) !important;
    color: var(--text-dark) !important;
    border: 1px solid rgba(0,0,0,0.3) !important;
    font-weight: 500;
}
.stTextInput label, .stNumberInput label, .stTextArea label,
.stDateInput label, .stTimeInput label, .stSelectbox label,
.stMultiSelect label, .stRadio label, .stCheckbox label {
    color: var(--text-light) !important;
    font-weight: 600;
}
.stSelectbox > div > div > div, .stMultiSelect > div > div > div {
    background: var(--bg-white) !important;
    color: var(--text-dark) !important;
    border: 1px solid rgba(0,0,0,0.3) !important;
}
[data-baseweb="select"] * { color: #111827 !important; }
[data-baseweb="menu"] { background: #FFFFFF !important; }
[data-baseweb="menu"] * { color: #111827 !important; background: #FFFFFF !important; }
[data-baseweb="popover"] { background: #FFFFFF !important; }
[data-baseweb="popover"] * { color: #111827 !important; }
ul[role="listbox"] { background: #FFFFFF !important; }
ul[role="listbox"] li { color: #111827 !important; }
ul[role="listbox"] li:hover { background: #F1F5F9 !important; }
.stDataFrame { background: var(--bg-white) !important; color: var(--text-dark) !important; }
.stDataFrame td, .stDataFrame th { color: var(--text-dark) !important; background: var(--bg-white) !important; }
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1E293B 0%, #0F172A 100%) !important;
}
section[data-testid="stSidebar"] *, section[data-testid="stSidebar"] label {
    color: var(--text-light) !important;
}
.dash-card, .rp-card, .metric-card {
    background: rgba(255,255,255,0.05);
    border: 1px solid rgba(255,255,255,0.1);
    border-radius: 16px; padding: 20px; margin-bottom: 16px;
}
.dash-card *, .rp-card *, .metric-card * { color: var(--text-primary) !important; }
.stButton > button {
    background: linear-gradient(135deg, var(--accent), var(--accent-hover));
    color: white !important; border: none;
    border-radius: 12px; padding: 10px 24px; font-weight: 600;
}
.status-pending    { color: var(--warning) !important; font-weight: 600; }
.status-ok         { color: var(--success) !important; font-weight: 600; }
.status-calibrated { color: var(--info)    !important; font-weight: 600; }
.status-installed  { color: var(--accent)  !important; font-weight: 600; }
.status-completed  { color: var(--success) !important; font-weight: 700; }
.status-rejected   { color: var(--error)   !important; font-weight: 600; }
[data-testid="stMetric"] {
    background: linear-gradient(135deg, rgba(59,130,246,0.3), rgba(96,165,250,0.2));
    border: 2px solid rgba(59,130,246,0.5); border-radius: 12px; padding: 15px;
}
[data-testid="stMetricValue"] { color: #60A5FA !important; }
[data-testid="stMetricLabel"] { color: #94A3B8 !important; }
"""

def inject_global_css():
    st.markdown(f"<style>{GLOBAL_CSS}</style>", unsafe_allow_html=True)

# =============================================================================
# AUDIT TRAIL
# =============================================================================
def log_audit(usuario, acao, tabela, registro_id, detalhes="", ip=""):
    try:
        return _gcs_append_row("logs_audit.csv", {
            "ID":          str(uuid.uuid4())[:8].upper(),
            "Data":        datetime.now().strftime("%d/%m/%Y"),
            "Hora":        datetime.now().strftime("%H:%M:%S"),
            "Usuario":     usuario, "Acao": acao, "Tabela": tabela,
            "Registro_ID": str(registro_id), "Detalhes": detalhes, "IP": ip
        })
    except Exception as e:
        logger.warning(f"log_audit erro: {e}")
        return False

def get_audit_logs(filtro_usuario=None, filtro_data=None, limite=100):
    try:
        logs = load_db("logs_audit.csv", [
            "ID","Data","Hora","Usuario","Acao","Tabela","Registro_ID","Detalhes","IP"
        ])
        if not logs.empty:
            if filtro_usuario: logs = logs[logs['Usuario'] == filtro_usuario]
            if filtro_data:    logs = logs[logs['Data']    == filtro_data]
            logs = logs.sort_values(['Data','Hora'], ascending=False)
            return logs.head(limite)
        return pd.DataFrame()
    except Exception as e:
        print(f"Erro ao obter logs: {e}")
        return pd.DataFrame()

# =============================================================================
# ASSINATURA DIGITAL
# =============================================================================
def gerar_hash_assinatura(usuario, tag, data, valor):
    texto = f"{usuario}|{tag}|{data}|{valor}"
    return hashlib.sha256(texto.encode()).hexdigest()[:16].upper()

def render_signature_pad(label, key_prefix):
    st.markdown(f"### {label}", unsafe_allow_html=True)
    signature = st.text_area(
        f"{label} — escreva o nome completo",
        key=f"{key_prefix}_nome",
        placeholder="Ex: João Silva",
        help="Assinatura digital"
    )
    if signature and len(signature.strip()) >= 3:
        fake_sig = base64.b64encode(
            f"SIGN:{signature.strip()}:{datetime.now().isoformat()}".encode()
        ).decode()
        return fake_sig
    return None

# =============================================================================
# NOTIFICAÇÕES
# =============================================================================
def criar_notificacao(destinatario, titulo, mensagem, tipo="info", acao_url=""):
    try:
        _gcs_append_row("notificacoes.csv", {
            "ID":           str(uuid.uuid4())[:8].upper(),
            "Data":         datetime.now().strftime("%d/%m/%Y"),
            "Hora":         datetime.now().strftime("%H:%M:%S"),
            "Destinatario": destinatario, "Titulo": titulo,
            "Mensagem":     mensagem, "Tipo": tipo,
            "Lida":         "Não", "Acao_URL": acao_url
        })
        inv("notificacoes.csv")
        return True
    except Exception as e:
        logger.warning(f"criar_notificacao erro: {e}")
        return False

def get_notificacoes(destinatario, apenas_nao_lidas=True, limite=50):
    try:
        notifs = load_db("notificacoes.csv", [
            "ID","Data","Hora","Destinatario","Titulo","Mensagem","Tipo","Lida","Acao_URL"
        ])
        if not notifs.empty:
            notifs = notifs[notifs['Destinatario'] == destinatario]
            if apenas_nao_lidas:
                notifs = notifs[notifs['Lida'] == "Não"]
            notifs = notifs.sort_values(['Data','Hora'], ascending=False)
            return notifs.head(limite)
        return pd.DataFrame()
    except:
        return pd.DataFrame()

def marcar_notificacao_lida(notif_id):
    try:
        notifs = load_db("notificacoes.csv", [
            "ID","Data","Hora","Destinatario","Titulo","Mensagem","Tipo","Lida","Acao_URL"
        ])
        if not notifs.empty:
            notifs.loc[notifs['ID'] == notif_id, 'Lida'] = "Sim"
            save_db(notifs, "notificacoes.csv")
            inv("notificacoes.csv")
            return True
        return False
    except:
        return False

def contar_notificacoes_nao_lidas(destinatario):
    try:
        notifs = get_notificacoes(destinatario, apenas_nao_lidas=True, limite=1000)
        return len(notifs)
    except:
        return 0

# =============================================================================
# FIX 3 — check_connection_status com cache TTL=30s
# Antes: chamada GCS em cada render do Admin
# Agora: resultado em cache durante 30 segundos
# =============================================================================
@st.cache_data(ttl=30, show_spinner=False)
def check_connection_status():
    """
    Verifica ligação ao GCS com cache de 30 segundos.
    Evita chamada de rede desnecessária em cada rerun.
    """
    try:
        client = _gcs_client()
        if client:
            bucket = client.bucket(GCS_BUCKET)
            bucket.reload()
            return True
        return False
    except:
        return False

def save_to_local_cache(key, data):
    try:
        cache_key = f"offline_cache_{key}"
        if cache_key not in st.session_state:
            st.session_state[cache_key] = []
        st.session_state[cache_key].append({
            "timestamp": datetime.now().isoformat(),
            "data": data.to_dict() if hasattr(data, 'to_dict') else data
        })
        if len(st.session_state[cache_key]) > 100:
            st.session_state[cache_key] = st.session_state[cache_key][-100:]
        return True
    except:
        return False

def get_from_local_cache(key, limite=50):
    try:
        cache_key = f"offline_cache_{key}"
        if cache_key in st.session_state:
            return st.session_state[cache_key][-limite:]
        return []
    except:
        return []

def add_action_to_queue(acao, dados, usuario):
    try:
        if "offline_action_queue" not in st.session_state:
            st.session_state["offline_action_queue"] = []
        st.session_state["offline_action_queue"].append({
            "id":        str(uuid.uuid4())[:8].upper(),
            "timestamp": datetime.now().isoformat(),
            "acao":      acao, "dados": dados,
            "usuario":   usuario, "estado": "pendente"
        })
        return True
    except:
        return False

def execute_offline_queue():
    resultados = {"sucessos": 0, "falhas": 0, "detalhes": []}
    try:
        if "offline_action_queue" not in st.session_state:
            return resultados
        queue = st.session_state["offline_action_queue"]
        for item in queue:
            if item["estado"] == "pendente":
                try:
                    if item["acao"] == "SAVE_DB":
                        df = pd.DataFrame(item["dados"]["data"])
                        save_db(df, item["dados"]["filename"])
                    item["estado"] = "executado"
                    resultados["sucessos"] += 1
                    resultados["detalhes"].append(f"✅ {item['acao']} - {item['id']}")
                except Exception as e:
                    item["estado"] = "falhou"
                    resultados["falhas"] += 1
                    resultados["detalhes"].append(f"❌ {item['acao']} - {item['id']} - {str(e)}")
        st.session_state["offline_action_queue"] = [
            i for i in queue if i["estado"] == "pendente"
        ]
        return resultados
    except:
        return resultados

def render_connection_indicator():
    st.markdown("""
    <script>
    function updateConnectionStatus() {
        const status = navigator.onLine ? 'online' : 'offline';
        const indicator = document.getElementById('connection-indicator');
        if (indicator) {
            indicator.className = `connection-status ${status}`;
            indicator.textContent = status === 'online' ? '🟢 Online' : '🔴 Offline';
        }
    }
    window.addEventListener('online',  updateConnectionStatus);
    window.addEventListener('offline', updateConnectionStatus);
    updateConnectionStatus();
    </script>
    <style>
    .connection-status {
        position:fixed; bottom:80px; right:20px; padding:8px 16px;
        border-radius:20px; font-weight:bold; font-size:0.85rem;
        z-index:9999; box-shadow:0 4px 6px rgba(0,0,0,0.3);
    }
    .connection-status.online  { background:#10B981; color:white; }
    .connection-status.offline { background:#EF4444; color:white; animation:pulse 2s infinite; }
    @keyframes pulse { 0%,100%{opacity:1} 50%{opacity:0.7} }
    </style>
    <div id="connection-indicator" class="connection-status online">🟢 Online</div>
    """, unsafe_allow_html=True)

def render_offline_banner():
    st.markdown("""
    <script>
    function checkOffline() {
        if (!navigator.onLine) {
            const b = document.getElementById('offline-banner');
            if (b) b.style.display = 'block';
        }
    }
    window.addEventListener('offline', checkOffline);
    window.addEventListener('online',  () => {
        const b = document.getElementById('offline-banner');
        if (b) b.style.display = 'none';
    });
    </script>
    <div id="offline-banner" style="display:none;background:#EF4444;color:white;
        padding:15px;border-radius:10px;margin-bottom:20px;text-align:center;">
        <strong>🔴 ESTÁ OFFLINE</strong> — Alterações guardadas localmente.
    </div>
    """, unsafe_allow_html=True)

def sync_data_when_online():
    if "offline_action_queue" not in st.session_state:
        return
    pendentes = [i for i in st.session_state["offline_action_queue"] if i["estado"] == "pendente"]
    if not pendentes:
        return
    if not check_connection_status():
        return
    resultados = execute_offline_queue()
    if resultados["sucessos"] > 0:
        st.success(f"✅ {resultados['sucessos']} ações sincronizadas!")
        _cached_load_db.clear()
    if resultados["falhas"] > 0:
        st.error(f"❌ {resultados['falhas']} ações falharam.")

# =============================================================================
# QR CODE
# =============================================================================
def gerar_qr_code_data(tag, obra, tipo_inst, url_base=""):
    qr_data = {
        "tag": tag, "obra": obra, "tipo": tipo_inst,
        "app": "GESTNOW", "v": "3.0",
        "url": f"{url_base}/instrumentacao?tag={tag}&obra={obra}" if url_base else f"/instrumentacao?tag={tag}"
    }
    return {
        "json":  json.dumps(qr_data),
        "b64":   base64.b64encode(json.dumps(qr_data).encode()).decode(),
        "short": f"GN|{tag}|{obra.replace(' ','_')}"
    }

def parse_qr_code_data(qr_string):
    try:
        if qr_string.startswith("{"):
            return json.loads(qr_string)
        if qr_string.startswith("ey"):
            decoded = base64.b64decode(qr_string).decode()
            return json.loads(decoded)
        if qr_string.startswith("GN|"):
            parts = qr_string.split("|")
            if len(parts) >= 3:
                return {"tag": parts[1], "obra": parts[2].replace("_"," "), "app":"GESTNOW"}
        return None
    except:
        return None

def render_qr_code_image(qr_data, size=200):
    import urllib.parse
    encoded = urllib.parse.quote(qr_data)
    return f"https://api.qrserver.com/v1/create-qr-code/?size={size}x{size}&data={encoded}"

def render_camera_scanner(label="Scan QR Code", key_prefix="qr_scan"):
    st.markdown(f"### 📱 {label}", unsafe_allow_html=True)
    uploaded = st.file_uploader("Upload de foto do QR Code",
        type=["png","jpg","jpeg"], key=f"{key_prefix}_upload")
    if uploaded:
        st.info("🔧 Leitura automática em desenvolvimento. Use o campo abaixo:")
    qr_manual = st.text_input("Dados do QR Code (formato: GN|TAG|OBRA ou JSON)",
        key=f"{key_prefix}_input")
    if qr_manual and len(qr_manual.strip()) > 5:
        return qr_manual.strip()
    return None

# =============================================================================
# SMTP
# =============================================================================
def get_smtp_config():
    smtp_server   = os.environ.get("SMTP_SERVER",    "")
    smtp_port     = os.environ.get("SMTP_PORT",      "587")
    smtp_user     = os.environ.get("SMTP_USER",      "")
    smtp_password = os.environ.get("SMTP_PASSWORD",  "")
    smtp_from_name= os.environ.get("SMTP_FROM_NAME", "GestNow")
    if smtp_server and smtp_user and smtp_password:
        return {
            "server": smtp_server, "port": int(smtp_port),
            "user": smtp_user, "password": smtp_password,
            "from_name": smtp_from_name, "from_email": smtp_user
        }
    return None

def enviar_email(destinatario, assunto, conteudo_html, conteudo_texto=""):
    import smtplib
    from email.mime.text      import MIMEText
    from email.mime.multipart import MIMEMultipart
    config = get_smtp_config()
    if not config:
        logging.warning("SMTP não configurado.")
        return False
    try:
        msg = MIMEMultipart("alternative")
        msg["Subject"] = f"[GestNow] {assunto}"
        msg["From"]    = f"{config['from_name']} <{config['from_email']}>"
        msg["To"]      = destinatario
        if conteudo_texto:
            msg.attach(MIMEText(conteudo_texto, "plain", "utf-8"))
        msg.attach(MIMEText(conteudo_html, "html", "utf-8"))
        server = smtplib.SMTP(config["server"], config["port"])
        server.starttls()
        server.login(config["user"], config["password"])
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        logging.error(f"Erro ao enviar email: {e}")
        return False

def get_email_template(tipo, dados=None):
    templates = {
        "validacao_horas": {
            "assunto": "Horas Validadas - {obra}",
            "html":    """<html><body style="font-family:Arial,sans-serif;padding:20px;">
                <h2>✅ Horas Validadas</h2>
                <p>Olá <strong>{tecnico}</strong>,</p>
                <p>As suas horas foram validadas com sucesso!</p>
                <p><strong>Obra:</strong> {obra}<br>
                <strong>Horas:</strong> {horas}h<br>
                <strong>Data:</strong> {data}<br>
                <strong>Validado por:</strong> {validador}</p>
                <p style="color:#94A3B8;font-size:0.85rem;">© GESTNOW v3.0</p>
            </body></html>""",
            "texto": "Horas validadas. Obra: {obra}, Horas: {horas}h"
        },
    }
    if tipo not in templates:
        return ("Notificação GestNow", "<p>Notificação do sistema GestNow</p>", "Notificação")
    template = templates[tipo]
    dados    = dados or {}
    assunto  = template["assunto"]
    html     = template["html"]
    texto    = template["texto"]
    for key, value in dados.items():
        assunto = assunto.replace("{" + key + "}", str(value))
        html    = html.replace("{" + key + "}", str(value))
        texto   = texto.replace("{" + key + "}", str(value))
    return (assunto, html, texto)

def notificar_por_email(destinatario_email, tipo_notificacao, dados):
    assunto, html, texto = get_email_template(tipo_notificacao, dados)
    return enviar_email(destinatario_email, assunto, html, texto)

def testar_smtp(email_teste):
    assunto, html, texto = get_email_template("validacao_horas", {
        "tecnico": "Teste", "obra": "Obra Teste", "horas": "8",
        "data": datetime.now().strftime("%d/%m/%Y"), "validador": "SYSTEM"
    })
    return enviar_email(email_teste, f"TESTE SMTP - {assunto}", html, texto)

# =============================================================================
# CLAUDE VISION — EXTRAÇÃO FOLHAS DE PONTO
# =============================================================================
def extrair_folha_ponto_vision(imagem_b64: str, media_type: str = "image/jpeg") -> dict:
    import anthropic, os, json

    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return {"sucesso": False, "dados": [], "raw": "", "erro": "API key não configurada."}

    prompt = """Analisa esta imagem de uma folha de ponto de obra.
Extrai TODOS os nomes de trabalhadores/técnicos e as horas que cada um trabalhou.

Devolve APENAS um JSON válido com este formato exacto, sem texto adicional:
{
  "trabalhadores": [
    {
      "nome": "Nome Completo do Trabalhador",
      "horas_total": 8.0,
      "dias": ["01/05/2025", "02/05/2025"]
    }
  ],
  "periodo": "período identificado na folha ou vazio",
  "obra": "nome da obra se visível ou vazio",
  "observacoes": "qualquer nota relevante ou vazio"
}

Regras:
- horas_total é a soma de todas as horas da semana/período desse trabalhador
- dias é a lista de datas em que o trabalhador aparece registado
- Se não conseguires ler um nome claramente, inclui o que consegues ver
- Se não houver horas numéricas, estima pelas marcas (X, /, tick)
- Responde APENAS com o JSON, sem markdown, sem ```"""

    try:
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type":       "base64",
                            "media_type": media_type,
                            "data":       imagem_b64,
                        }
                    },
                    {
                        "type": "text",
                        "text": prompt
                    }
                ]
            }]
        )

        raw_text = response.content[0].text.strip()
        raw_text = raw_text.replace("```json","").replace("```","").strip()
        dados_json = json.loads(raw_text)
        trabalhadores = dados_json.get("trabalhadores", [])

        return {
            "sucesso":  True,
            "dados":    trabalhadores,
            "periodo":  dados_json.get("periodo",""),
            "obra":     dados_json.get("obra",""),
            "obs":      dados_json.get("observacoes",""),
            "raw":      raw_text,
            "erro":     ""
        }

    except json.JSONDecodeError as e:
        return {
            "sucesso": False, "dados": [],
            "raw": raw_text if 'raw_text' in dir() else "",
            "erro": f"Erro ao interpretar JSON: {e}"
        }
    except Exception as e:
        return {
            "sucesso": False, "dados": [], "raw": "",
            "erro": f"Erro Vision API: {e}"
        }
